import streamlit as st
import pdfplumber  
import docx2txt  
import google.generativeai as genai
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

# Configure Gemini API
genai.configure(api_key="AIzaSyCOjWzY0tvaeMxjHK1mqDLX-EeN3h-VDLA")  
model = genai.GenerativeModel(model_name="gemini-2.5-pro")

st.title("AI-Based Smart Exam Question Generator 📝🤖")

# Upload College Banner
uploaded_logo = st.file_uploader("Upload College Logo", type=["png", "jpg", "jpeg"])

# File uploader
uploaded_file = st.file_uploader("Upload Syllabus (PDF or DOCX)", type=["pdf", "docx"])

def extract_text_from_pdf(uploaded_pdf):
    text = ""
    try:
        with pdfplumber.open(uploaded_pdf) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n" if page.extract_text() else ""
    except Exception as e:
        st.error(f"Error extracting PDF text: {e}")
    return text.strip()

def extract_text_from_docx(uploaded_docx):
    try:
        text = docx2txt.process(uploaded_docx)
    except Exception as e:
        st.error(f"Error extracting DOCX text: {e}")
        text = ""
    return text.strip()

if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    if file_type == "pdf":
        syllabus_content = extract_text_from_pdf(uploaded_file)
    elif file_type == "docx":
        syllabus_content = extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file format. Please upload a PDF or DOCX file.")
        syllabus_content = None

    if not syllabus_content.strip():
        st.error("Syllabus content is required to generate questions. Please upload a valid syllabus file.")
    else:
        st.subheader("Configure Question Paper Format")

        total_marks = st.number_input("Total Marks", min_value=25, max_value=100, value=25, step=5)
        duration = st.selectbox("Exam Duration", ["45 Minutes", "1.5 Hours", "3 Hours"])
        college_name = st.text_input("Enter College Name")
        course_code = st.text_input("Enter Course Code")
        course_name = st.text_input("Enter Course Name")

        sections = st.multiselect("Select Sections", ["PART A", "PART B", "PART C"], default=["PART A"])
        
        section_details = {}
        
        for section in sections:
            st.subheader(f"Configure {section}")
            section_details[section] = {
                "total_questions": st.number_input(f"Total number of questions in {section}", min_value=1, max_value=100, value=5, step=1),
                "marks_per_question": st.number_input(f"Marks per question in {section}", min_value=1, max_value=total_marks, value=5, step=1),
                "question_type": st.selectbox(f"Select Question Type for {section}", ["MCQ", "Short Answer", "Long Answer"]),
                "k_level": st.selectbox(f"K-Level for {section}", ["K1 - Remember", "K2 - Understand", "K3 - Apply", "K4 - Analyze", "K5 - Evaluate", "K6 - Create"]),
                "unit_distribution": {}
            }
            
            st.subheader(f"Select Units for {section}")
            selected_units = st.multiselect(f"Choose units to include in {section}", ["Unit 1", "Unit 2", "Unit 3", "Unit 4", "Unit 5"])
            
            remaining_questions = section_details[section]["total_questions"]
            for unit in selected_units:
                max_q_per_unit = min(remaining_questions, 100)
                section_details[section]["unit_distribution"][unit] = st.number_input(f"Number of questions from {unit} in {section}", min_value=0, max_value=max_q_per_unit, value=min(remaining_questions, 5), step=1)
                remaining_questions -= section_details[section]["unit_distribution"][unit]
        
        allocated_marks = sum(section_details[section]["marks_per_question"] * section_details[section]["total_questions"] for section in sections)
        
        if allocated_marks != total_marks:
            st.error(f"Total allocated marks ({allocated_marks}) do not match the specified total marks ({total_marks}). Adjust question counts or marks per question.")
        else:
            if st.button("Generate Question Paper"):
                prompt = f"""
You are an AI-powered question paper generator. Ensure the generated exam paper strictly adheres to the syllabus below:

### Syllabus:
{syllabus_content}

### Exam Details:
- College Name: {college_name}
- Course Code: {course_code}
- Course Name: {course_name}
- Total Marks: {total_marks}
- Exam Duration: {duration}

### Question Paper Structure:
"""
                for section, details in section_details.items():
                    prompt += f"""
#### {section}
- Question Type: {details['question_type']}
- K-Level: {details['k_level']}
- Total Questions: {details['total_questions']}
- Marks per Question: {details['marks_per_question']}

Generate {details['total_questions']} questions for {section}, ensuring:
- They strictly adhere to the K-Level: {details['k_level']}.
- They are based only on the provided syllabus.
- The specified number of questions come from each unit, but do not explicitly mention the unit in the question paper.
"""
                response = model.generate_content([prompt], generation_config={"temperature": 0.3, "top_p": 0.8})
                question_paper = response.text if response else "Error generating questions."
                
                doc = Document()
                if uploaded_logo:
                    image_stream = io.BytesIO(uploaded_logo.read())
                    doc.add_picture(image_stream, width=Inches(5.0))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                doc.add_paragraph(question_paper)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button("Download as DOCX", buffer, "question_paper.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")



